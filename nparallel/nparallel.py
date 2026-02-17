import concurrent.futures
import os
import tempfile
from threading import Lock
import argparse
import xml.etree.ElementTree as ET
from datetime import datetime

import subprocess
import ipaddress
import hashlib
import socket
import shutil
import sys
from itertools import repeat


import operator ## for grouping operation
import itertools ## for grouping operation
import collections ## for creating a defaultdict

class NmapCommand: 
    cmd_str = ""

    def __init__(self, cmd_str:str):
        self.cmd_str = cmd_str
        
    def get_id (self):
        """Get a unique command ID based on the used nmap command.
        Returns:
            A unique command ID.
        """        
        # sort nmap commands, to calculate same hash value, even if only 
        # the place of single parameters have changed.
        cmd_list_sorted = self.cmd_str.split(" ")
        cmd_list_sorted.sort()
        cmd_list_sorted_str = ' '.join(cmd_list_sorted)

        # calculate hash_value and trunace
        command_id = hashlib.sha256(cmd_list_sorted_str.encode()).hexdigest()[:8]
        return command_id

        

class NmapScan:
    ip = None
    nmap_cmd = None
    result = None

    def __init__(self, ip, nmap_cmd:NmapCommand):
        self.ip = ip
        self.nmap_cmd = nmap_cmd

    def get_id (self):
        return f"{str(self.ip)}_{self.nmap_cmd.get_id()}"

    def run (self, scan_cache_path):
        scan_id = self.get_id()
        # craft command
        command = "{} -oX {} -oN {} -oG {} {}".format(
            self.nmap_cmd.cmd_str,
            os.path.join(scan_cache_path, scan_id + ".xml"),
            os.path.join(scan_cache_path, scan_id + ".txt"),
            os.path.join(scan_cache_path, scan_id + ".grep"),
            self.ip,
        )
        # run command     
        with open(os.path.join(scan_cache_path, scan_id + ".log"), 'w') as log_file:
            self.result = subprocess.run(command.split(), shell=False, stdout=log_file, stderr=log_file)


class NmapExportBuilder:
    def export (self, file_ext, cache_dir, nmap_cmd, ip_list, start_time, end_time, out_file_path):
        "Export a file_type file"

        src_file_ext = file_ext
        build_and_export_method = None
        
        match file_ext:
            case ".xml":
                build_and_export_method = self._build_and_export_xml
            case ".txt":
                build_and_export_method = self._build_and_export_normal
            case ".grep":
                build_and_export_method = self._build_and_export_grepable
            case ".log":
                build_and_export_method = self._build_and_export_log
            case ".csv":
                src_file_ext = ".xml"
                build_and_export_method = self._build_and_export_csv
            case ".xlsx":
                src_file_ext = ".xml"
                build_and_export_method = self._build_and_export_xlsx 
            case ".docx":
                src_file_ext = ".xml"
                build_and_export_method = self._build_and_export_docx          
            case _:
                print (f"[!] Unknown file extension {file_ext}. Export aborted")
                exit(1)
        
        # get source files
        files = self.get_files(src_file_ext, cache_dir, ip_list)
        # get raw file content
        build_and_export_method(files, nmap_cmd, start_time, end_time, out_file_path)
    

    def get_files (self, file_ext, cache_dir, ip_list):
        "Get list of files with file_ext from cache_dir"
        files = set()
        # look for all files ending with file_ext
        for f in os.listdir(cache_dir):
            # check if file extension matches and IP address part of the filename is in ip_list
            if f.endswith(file_ext) and ipaddress.IPv4Address(f.split('_')[0]) in ip_list:
                file_path = os.path.join(cache_dir, f)
                files.add(file_path)
        return files    


    def _build_and_export_xml (self, files:set, nmap_cmd:NmapCommand, start_time:datetime, end_time:datetime, out_file_path:str):
        "Merge all files in set to one single XML file and return content as string"
        out_str = ""        

        # add header
        with open (list(files)[0], 'r', encoding='utf-8') as xml_file:        
            xml = ET.parse(xml_file)
            out_str += '<?xml version="1.0" encoding="UTF-8"?>\n'
            out_str += '<!DOCTYPE nmaprun>\n'
            out_str += '<?xml-stylesheet href="file:///usr/share/nmap/nmap.xsl" type="text/xsl"?>\n'
            out_str += '<!-- Nmap result merged with nparallel.py https://github.com/ronnydo/nparallel -->\n'
            out_str += '<nmaprun scanner="nmap" args="{}" start="{}" startstr="{}" version="7.80" xmloutputversion="1.04">\n'.format(
                nmap_cmd.cmd_str,
                int(start_time.timestamp()), # Linux timestamp
                start_time.strftime("%a %b %d %H:%M:%S %Y"), #Fri Mar 21 22:11:16 2025
            )
            scaninfo_tag = xml.find('scaninfo')
            out_str += ET.tostring(scaninfo_tag, encoding='unicode', method='xml')
            verbose_tag = xml.find('verbose')
            out_str += ET.tostring(verbose_tag, encoding='unicode', method='xml')
            debugging_tag = xml.find('debugging')
            out_str += ET.tostring(debugging_tag, encoding='unicode', method='xml')
        
        # add hosts
        host_number = 0
        for xml_file_path in files:
            with open (xml_file_path, 'r', encoding='utf-8') as xml_file:
                xml = ET.parse(xml_file)
                for host in xml.findall('host'):
                    out_str += ET.tostring(host, encoding='unicode', method='xml')
                    host_number += 1
        
        # add footer
        out_str += ('<runstats><finished time="{}" timestr="{}" elapsed="{}" summary="Nmap done at {}; ' + str(host_number) + ' IP address scanned in {} seconds" exit="success"/>\n').format(
            int(end_time.timestamp()), # Linux timestamp
            end_time.strftime("%a %b %d %H:%M:%S %Y"),
            "{:.2f}".format((end_time - start_time).total_seconds()),
            end_time.strftime("%a %b %d %H:%M:%S %Y"),
            "{:.2f}".format((end_time - start_time).total_seconds()),
        )
        out_str += '</runstats>\n'
        out_str += '</nmaprun>'

        # write file
        with open(out_file_path, "w") as out_file:
            out_file.writelines(out_str)


    def _build_and_export_normal (self, files:set, nmap_cmd:NmapCommand, start_time:datetime, end_time:datetime, out_file_path:str):
        "Merge all files in set to one single txt file and return content as string"
        out_str = ""        

        # add header
        out_str += f'# Nmap scan initiated {start_time.strftime("%a %b %d %H:%M:%S %Y")} as: {nmap_cmd.cmd_str}\n'
        
        # add hosts
        hosts_total = len(files)
        hosts_up = len(files)
        for file_path in files:
            with open (file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                # remove first line
                lines.pop(0)
                # read all lines from first and second block
                block = 1                    
                for line in lines:
                    out_str += line                       
                    if line.startswith ("Nmap scan report for") and line.strip().endswith("[host down]"):
                        # decrease hosts_up of results are empty
                        hosts_up -= 1                     
                    if line.strip() == "":
                        block += 1
                    if block == 3:
                        break
        
        # add footer
        out_str += "# Nmap done at {} -- {} IP addresses ({} hosts up) scanned in {} seconds".format(
            end_time.strftime("%a %b %d %H:%M:%S %Y"),
            hosts_total,
            hosts_up,           
            "{:.2f}".format((end_time - start_time).total_seconds()),
        )

        # write file
        with open(out_file_path, "w") as out_file:
            out_file.writelines(out_str)
    

    def _build_and_export_grepable (self, files:set, nmap_cmd:NmapCommand, start_time:datetime, end_time:datetime, out_file_path:str):
        "Merge all files in set to one single grepable file and return content as string"
        out_str = ""        

        # add header
        with open (list(files)[0], 'r', encoding='utf-8') as file:
            # read first line but set line manually
            file.readline()
            out_str += f'# Nmap scan initiated {start_time.strftime("%a %b %d %H:%M:%S %Y")} as: {nmap_cmd.cmd_str}\n'
            # read second line with scanned ports information
            out_str += file.readline()            
        
        # add hosts
        hosts_total = len(files)
        hosts_up = len(files)
        for file_path in files:
            with open (file_path, 'r', encoding='utf-8') as file:
                for line in file.readlines():
                    if line.startswith("# ") == False:
                        out_str += line
                        # decrease hosts_up of host is down
                        if line.strip().endswith("Status: Down"):
                            hosts_up -= 1
        
        # add footer
        out_str += "# Nmap done at {} -- {} IP addresses ({} hosts up) scanned in {} seconds".format(
            end_time.strftime("%a %b %d %H:%M:%S %Y"),
            hosts_total,
            hosts_up,           
            "{:.2f}".format((end_time - start_time).total_seconds()),
        )

        # write file
        with open(out_file_path, "w") as out_file:
            out_file.writelines(out_str)


    def _build_and_export_log (self, files:set, nmap_cmd:NmapCommand, start_time:datetime, end_time:datetime, out_file_path:str):
        "Merge all files in set to one single log file and return content as string"        
        out_str = ""        

        # add header
        out_str += f'Starting Nmap ( https://nmap.org ) at {start_time.strftime("%a %b %d %H:%M:%S %Y")}\n'
        
        # add hosts
        hosts_total = len(files)
        hosts_up = len(files)
        for file_path in files:
            with open (file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                # remove first line
                lines.pop(0)
                # read all lines from first and second block
                block = 1                    
                for line in lines:
                    out_str += line                        
                    if line.startswith ("Nmap scan report for") and line.strip().endswith("[host down]"):
                        # decrease hosts_up of results are empty
                        hosts_up -= 1
                    if line.strip() == "":
                        block += 1
                    if block == 3:
                        break
        
        # add footer
        out_str += 'Read data files from: /usr/bin/../share/nmap\n'
        out_str += 'Nmap done: {} IP addresses ({} hosts up) scanned in {} seconds'.format(
            hosts_total,
            hosts_up,           
            "{:.2f}".format((end_time - start_time).total_seconds()),
        )

        # write file
        with open(out_file_path, "w") as out_file:
            out_file.writelines(out_str)


    def _build_and_export_csv (self, files:set, nmap_cmd:NmapCommand, start_time:datetime, end_time:datetime, out_file_path:str):
        "Merge all files in set to one single CSV file and return content as string"
        out_str = ""        

        # add header
        out_str += "IP address\tHostnames\tPort\tProto\tService\tInfo\n"
  
        # add body 
        for xml_file_path in files:
            ip = ""
            hostnames = set()
            with open (xml_file_path, 'r', encoding='utf-8') as xml_file:
                xml = ET.parse(xml_file)
                host = xml.find('host')

                # only hosts that are up
                if host.find('status').get('state') == 'up':
                    ip_address = host.find("address").get("addr")
                    
                    for hostname in host.find('hostnames'):
                        hostnames.add(hostname.get("name"))

                    for port_tag in host.find('ports').findall('port'):
                        state_tag = port_tag.find('state')
                        if state_tag.get('state') == 'open':       
                            proto = port_tag.get('protocol')                 
                            port = port_tag.get('portid')

                            service_tag = port_tag.find('service')
                            service = ""
                            info = ""
                            if service_tag != None:
                                service = service_tag.get("name") if service_tag.get("name") else ""
                                info = service_tag.get("product") if service_tag.get("product") else ""
                                info += f" {service_tag.get('version')}" if service_tag.get("version") else ""
                                info += f" ({service_tag.get('extrainfo')})" if service_tag.get("extrainfo") else ""

                            # add entry
                            out_str += f"{ip_address}\t{'/'.join(sorted(hostnames))}\t{port}\t{proto}\t{service}\t{info}\n"        
        # write file
        with open(out_file_path, "w") as out_file:
            out_file.writelines(out_str)


    def _build_and_export_xlsx (self, files:set, nmap_cmd:NmapCommand, start_time:datetime, end_time:datetime, out_file_path:str):
        "Merge all files in set to one single XSLX file and return content as string"        
            
        workbook = xlsxwriter.Workbook(out_file_path)       
        worksheet = workbook.add_worksheet()

        # add (dummy) header line        
        worksheet.write_row("A1", ())

        # get data
        row = 1   
        for xml_file_path in files:
            ip = ""
            hostnames = set()
            with open (xml_file_path, 'r', encoding='utf-8') as xml_file:
                xml = ET.parse(xml_file)
                host = xml.find('host')

                # only hosts that are up
                if host.find('status').get('state') == 'up':
                    ip_address = host.find("address").get("addr")
                    
                    for hostname in host.find('hostnames'):
                        hostnames.add(hostname.get("name"))

                    for port_tag in host.find('ports').findall('port'):
                        state_tag = port_tag.find('state')
                        if state_tag.get('state') == 'open':       
                            proto = port_tag.get('protocol')                 
                            port = port_tag.get('portid')

                            service_tag = port_tag.find('service')
                            service = ""
                            info = ""
                            if service_tag != None:
                                service = service_tag.get("name") if service_tag.get("name") else ""
                                info = service_tag.get("product") if service_tag.get("product") else ""
                                info += f" {service_tag.get('version')}" if service_tag.get("version") else ""
                                info += f" ({service_tag.get('extrainfo')})" if service_tag.get("extrainfo") else ""

                            # add entry
                            worksheet.write_row(row, 0, (ip_address, '/'.join(sorted(hostnames)), port, proto, service, info))
                            row += 1

        # add real header
        worksheet.add_table(f"A1:F{row})", {'columns': [
            {'header': 'IP address'},
            {'header': 'Hostname'},
            {'header': 'Port'},
            {'header': 'Proto'},
            {'header': 'Service'},
            {'header': 'Info'},
        ]})
        
        # save doc
        workbook.close()


    def _build_and_export_docx (self, files:set, nmap_cmd:NmapCommand, start_time:datetime, end_time:datetime, out_file_path:str):
        "Merge all files in set to one single XSLX file and return content as string"      

        from docx import Document
        from docx.shared import Pt  
        from docx.enum.section import WD_ORIENT,WD_SECTION
            
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri Light'
        font.size = Pt(10)
        
        # switch to landscape mode
        landscape = False
        if landscape:
            section = document.sections[-1]
            # switch width and height
            new_width, new_height = section.page_height, section.page_width        
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height

        # add header        
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "IP"
        table.rows[0].cells[1].text = "Hostnames"
        table.rows[0].cells[2].text = "Ports"
        table.rows[0].cells[3].text = "Service"
        table.rows[0].cells[4].text = "Info"

        hosts_dict = {}

        # read data from files
        for xml_file_path in files:
            with open (xml_file_path, 'r', encoding='utf-8') as xml_file:
                xml = ET.parse(xml_file)
                host = xml.find('host')                

                # only hosts that are up
                if host.find('status').get('state') == 'up':  
                    hostnames = set()                               
                    ip = ""
                    port_proto_dict = {}
                    service_dict = {}
                    info_dict = {}

                    ip_address = ipaddress.IPv4Address(host.find("address").get("addr"))
                    
                    for hostname in host.find('hostnames'):
                        hostnames.add(hostname.get("name"))


                    for port_tag in host.find('ports').findall('port'):
                        state_tag = port_tag.find('state')
                        if state_tag.get('state') == 'open':   
                            port_proto = f"{port_tag.get('portid')}/{port_tag.get('protocol')}"
                            port = port_tag.get('portid')
                            proto = port_tag.get('protocol')

                            service_tag = port_tag.find('service')
                            service = ""
                            info = ""
                            if service_tag != None:
                                service = service_tag.get("name") if service_tag.get("name") else ""
                                info = service_tag.get("product") if service_tag.get("product") else ""
                                info += f" {service_tag.get('version')}" if service_tag.get("version") else ""
                                info += f" ({service_tag.get('extrainfo')})" if service_tag.get("extrainfo") else ""
                            
                            port_proto_dict[f"{port_tag.get('protocol')}_{port_tag.get('portid')}"] = f"{port_tag.get('portid')}/{port_tag.get('protocol')}"
                            service_dict[f"{port_tag.get('protocol')}_{port_tag.get('portid')}"] = service
                            info_dict[f"{port_tag.get('protocol')}_{port_tag.get('portid')}"] = info
                    
                    hosts_dict[ip_address] = {
                        'ip_address': ip_address,
                        'hostnames_list': sorted(hostnames),
                        'port_proto_list': port_proto_dict.values(),
                        'services_list': service_dict.values(),
                        'info_list': info_dict.values()
                    }
        
        # add body   
        for host_ip in sorted(hosts_dict.keys()):               
            row = table.add_row()
            host_values = hosts_dict.get(host_ip)
            row.cells[0].text = str(host_values.get('ip_address'))
            row.cells[1].text = '\n'.join(sorted(host_values.get('hostnames_list')))
            row.cells[2].text = '\n'.join(host_values.get('port_proto_list'))
            row.cells[3].text = '\n'.join(host_values.get('services_list'))
            row.cells[4].text = '\n'.join(host_values.get('info_list'))
                   
        # save doc        
        document.save(out_file_path)



class Nparallel:
    CACHE_DIR = '.cache'
    COMMAND_FILE_PREFIX = "_nmap_command_"

    def __init__(self):
        self.args = self.get_args()

    def run_nmap(self, scan:NmapScan, tmp_dir):
        scan.run(tmp_dir)
        return scan

    def get_args(self):
        "Parse and return command line arguments"
        parser = argparse.ArgumentParser(
            prog="Nparallel",
            description="Speed up Nmap scans by running them in parallel",
#            usage="\n Run Nmap scan:\n" \
#                  " > nparallel nmap [nmap-args] -iL FILE \n"\
#                  " > nparallel nmap -v --top-ports 100 -oX results.xml -iL targets.txt "
        )

        sub_parsers = parser.add_subparsers(help='command', dest="command")        

        parser_run = sub_parsers.add_parser('nmap', 
            help='Run scan',                                            
            usage="\n" \
                " > nparallel nmap [args below / arbitrary nmap args] -iL TARGETS_FILE\n" \
                " > nparallel nmap -v --top-ports 100 -oX results.xml -iL targets.txt")
        
        # nparallel nmap ...
        parser_run.add_argument('-iL', '--input-list', required=True, help='Input from list of hosts/networks')
        parser_run.add_argument('--force-scan', action='store_true', help='Scan hosts even if cache entry exists. Updates cache.')
        parser_run.add_argument('-t', '--threads', type=int, default=100, help="Number of parallel nmap threads")
        parser_run.add_argument('-oA', '--out-all', default=None, help='Output in the three major formats at once (normal, XML and grepable)')
        parser_run.add_argument('-oN', '--out-normal', default=None, help='Output scan in Normal format')
        parser_run.add_argument('-oG', '--out-grepable', default=None, help='Output scan in Grepable format')
        parser_run.add_argument('-oX', '--out-xml', default=None, help='Output scan in XML')
        parser_run.add_argument('-oL', '--out-log', default=None, help='Output scan in Log format')
        parser_run.add_argument('-oC', '--out-csv', default=None, help='Output hosts with open ports in CSV format')
        parser_run.add_argument('-oE', '--out-excel', default=None, help='Output hosts with open ports in XLSX format')
        parser_run.add_argument('-oW', '--out-word', default=None, help='Output hosts with open ports in DOCX format')

        # ls
        parser_ls = sub_parsers.add_parser('ls', help='Show cache')    
        parser_ls.add_argument('cmd_id', nargs='*', help='Inspect results of <cmd_id>. Merges results if multiple cmd_ids are provided.')
        parser_ls.add_argument('-g', '--grouped', action='store_true', help='Group IP addresses and ports for better readability')
        
        # rm
        parser_rm = sub_parsers.add_parser('rm', help='Remove from cache')
        parser_rm.add_argument('cmd_id', nargs='*', help='Nmap command ID(s) of the entries to be deleted')
        

        # interpret remaining arguments as nmap args
        args, unknown = parser.parse_known_args()
        args.nmap_command = f"nmap {' '.join(unknown)}"        

        return args       
    

    def resolve_to_ip_addresses(self, targets_file): # -> list[ipaddress.IPv4Address]:
        """Resolve hosts, which are in CIDR notation or hostnames, to IP addresses. Removes doubling.
        Args:
            targets_file: The file containing the target hosts in CIDR notation.
        Returns:
            An ordererd list of IP addresses.
        """
        resolved_ips = set()
        with open(os.path.expanduser(targets_file), 'r') as file:
            for line in file:
                line = line.strip()
                if line:
                    try:
                        network = ipaddress.ip_network(line, strict=False)
                        for ip in network.hosts():
                            resolved_ips.add(ip)
                    except ValueError:
                        try: 
                            # try resolving string if it isn't an IP
                            resolved_ip = socket.gethostbyname(line)
                            resolved_ips.add(ipaddress.ip_address(resolved_ip))
                        except:
                            print(f"Invalid network address '{line}' in '{targets_file}")
                            print(f"QUITTING!")
                            exit (1)
        return sorted(resolved_ips)
    
    def get_scaninfo_path (self, cmd_id:str):
        return os.path.join (self.CACHE_DIR, cmd_id, f"{self.COMMAND_FILE_PREFIX}{cmd_id}")
    
    def get_scan_cache_path (self, cmd_id:str):
        return os.path.join (self.CACHE_DIR, cmd_id)

    def get_nmap_base_cmd(self, cmd_id:str):  
        "reads nmap base command from scaninfo file. Returns None if scan_id does not exist"  
        nmap_base_cmd = None
        scaninfo_path = self.get_scaninfo_path(cmd_id)
        if os.path.exists(scaninfo_path):         
            with open(scaninfo_path, 'r') as scaninfo_file:
                nmap_base_cmd = scaninfo_file.readline()
        return nmap_base_cmd

    def get_scans(self, cmd_id, resolved_ips, force_scan):
        "Get lists of finished and unfinished scans based on resolved_ips"
        scans = []
        scans_open = []
        scans_finished = []

        # use scans from cache, if force_scan is false
        if force_scan == False:
            scans_finished = [f for f in os.listdir(self.get_scan_cache_path(cmd_id)) if f.endswith(".xml")]
        for ip in resolved_ips:
            nmap_scan = NmapScan(
                ip = ip,
                nmap_cmd = NmapCommand(self.get_nmap_base_cmd(cmd_id))
            )    
            if f"{nmap_scan.get_id()}.xml" not in scans_finished:
                scans_open.append(nmap_scan)
            scans.append(nmap_scan)
        return scans, scans_open
    
    def get_finished_scans (self, cmd_id):
        "Return XML file paths where scan is finished"
        scan_cache_path = self.get_scan_cache_path(cmd_id)
        scans_finished = [f for f in os.listdir(scan_cache_path) if f.endswith(".xml")]
        return scans_finished
      
    
    def init_scan_cache(self):
        "Init cache directory for specific command"
        nmap_cmd = NmapCommand(self.args.nmap_command)
        cache_dir = os.path.join(os.getcwd(), self.CACHE_DIR, nmap_cmd.get_id())

        # init scan folder
        if not os.path.exists(cache_dir):
            os.makedirs(cache_dir)
        
        # create scan info file
        scaninfo_file_path = os.path.join (cache_dir, f"{self.COMMAND_FILE_PREFIX}{nmap_cmd.get_id()}")
        if not os.path.exists(scaninfo_file_path):         
            with open(scaninfo_file_path, 'w') as scaninfo_file:
                scaninfo_file.write(nmap_cmd.cmd_str)

        return cache_dir


    def get_cache_info(self):
        scans = {} 
        base_cache_dir = os.path.join(os.getcwd(), self.CACHE_DIR)
        if (os.path.exists(base_cache_dir)):
            for cmd_id in os.listdir(base_cache_dir):
                scaninfo_path = self.get_scaninfo_path(cmd_id)
                            
                ip_adresses = list()
                for scan in self.get_finished_scans(cmd_id):
                    ip = ipaddress.IPv4Address(scan.split("_")[0])
                    ip_adresses.append(ip)

                if os.path.exists(scaninfo_path):         
                    with open(scaninfo_path, 'r') as scaninfo_file:
                        scans[cmd_id] = {
                            "cmd_id": cmd_id,
                            "nmap_base_cmd": self.get_nmap_base_cmd(cmd_id),
                            "scans_finished": self.get_finished_scans(cmd_id),
                            "scan_groups": self.group_ip_addresses(ip_adresses)
                        }
        return scans
    
    
    def get_open_ports (self, cmd_id, scanfile_xml_name):
        scan_cache = self.get_scan_cache_path(cmd_id)
        "Get list of open TCP and udp ports from nmap XML file"
        ports_tcp = set()
        ports_udp = set()
        with open(os.path.join(scan_cache, scanfile_xml_name), 'r') as xml_file:
            xml = ET.parse(xml_file)
            # should actually be just one host per file
            for host in xml.findall('host'):
                for port in host.findall('ports/port'):
                    # ensure port state is up
                    state = port.find("state")
                    if state.get("state") == "open":
                        # add tcp
                        if port.get("protocol") == "tcp":
                            ports_tcp.add(port.get("portid"))
                        # add udp
                        else:
                            ports_udp.add(port.get("portid"))
        return ports_tcp, ports_udp



    def remove_cache_entries(self, cmd_ids=[]):
        "Remove single or all cache entires"
        if cmd_ids and len(cmd_ids) > 0:
            for cmd_id in cmd_ids:
                scan_cache = self.get_scan_cache_path(cmd_id)
                if os.path.exists(scan_cache):
                    shutil.rmtree(scan_cache)
                    print (f"[+] Entry with cmd_id '{cmd_id}' removed")
                else:
                    print (f"[!] Entry for cmd_id '{cmd_id}' not found")    
        else:
            if input("[?] Delete entire cache? [y/N]") == "y":
                shutil.rmtree(self.CACHE_DIR)
                print ("[+] Cache cleared.")


    def group_ip_addresses(self, ip_addresses):
        groups = []
        for _, g in itertools.groupby(enumerate(sorted(ip_addresses)), lambda ix: ix[0] - int(ix[1])):
            group = list(map(operator.itemgetter(1), g))
            if len(group) > 1:
                groups.append(f"{group[0]}-{str(group[-1]).split('.')[-1]}")
            else:
                groups.append(str(group[0]))
        return groups

    def group_ports(self, ports):
        groups = []
        for _, g in itertools.groupby(enumerate(sorted(ports, key=int)), lambda ix: ix[0] - int(ix[1])):
            group = list(map(operator.itemgetter(1), g))
            if len(group) > 1:
                groups.append(f"{group[0]}-{str(group[-1]).split('.')[-1]}")
            else:
                groups.append(str(group[0]))
        return groups
    


    def print_cache_info(self, scans):
        if len (scans) > 0: 
            print (f"[*] Cache contains {len(scans)} scans:\n")  
            print (f"Cmd id     Finished\tNmap base command")
            print (f"---        ---     \t---")
            for entry in scans.values():
                    print (f"{entry['cmd_id']}   {len(entry['scans_finished'])}    \t{entry['nmap_base_cmd']}")
        else:
            print ("[*] Cache is empty")            


    def get_cmd_info (self, cmd_ids):
        cmd_ids_info = list()
        hosts_finished = set()
        hosts_with_open_ports = set()
        ports_tcp_open = set()
        ports_udp_open = set()

        for cmd_id in cmd_ids:

            scaninfo_path = self.get_scaninfo_path(cmd_id)

            cmd_info = None

            if (os.path.exists(scaninfo_path)):
                cmd_ids_info.append({
                    "cmd_id": cmd_id,
                    "nmap_base_cmd": self.get_nmap_base_cmd(cmd_id)
                })
                
                for scanfile_xml_name in self.get_finished_scans(cmd_id):          
                    ip_str = scanfile_xml_name.split("_")[0]
                    hosts_finished.update([ip_str])
                    ports_tcp, ports_udp = self.get_open_ports (cmd_id, scanfile_xml_name)
                    if len(ports_tcp) > 0 or len(ports_udp) > 0:
                        hosts_with_open_ports.update([ip_str])
                        ports_tcp_open.update(ports_tcp)
                        ports_udp_open.update(ports_udp)

            else:
                print (f"[!] Unkown cmd id '{cmd_id}'")        

        cmd_info = {
            "cmd_ids_info": cmd_ids_info,
            "hosts_finished": sorted(ipaddress.IPv4Address(ip_str) for ip_str in hosts_finished), # convert to IPv4 to sort correctly
            "hosts_with_open_ports": sorted(ipaddress.IPv4Address(ip_str) for ip_str in hosts_with_open_ports), # convert to IPv4 to sort correctly
            "ports_tcp_open": sorted(ports_tcp_open, key=int),
            "ports_udp_open": sorted(ports_udp_open, key=int),
        }

        return cmd_info


    def print_cmd_info(self, cmd_info, grouped):
        if cmd_info: 
            print (f"[*] Nmap base command(s):")
            print (f"{'\n'.join(x.get("cmd_id")+": "+x.get("nmap_base_cmd") for x in cmd_info['cmd_ids_info'])}")

            print (f"\n[+] Hosts finished (\033[92m{len(cmd_info['hosts_finished'])}\033[0m):")
            if grouped:
                print (f"{' '.join(self.group_ip_addresses(cmd_info['hosts_finished']))}")
            else:
                print (f"{' '.join([str(x) for x in cmd_info['hosts_finished']])}")

            print (f"\n[+] Hosts with open ports (\033[92m{len(cmd_info['hosts_with_open_ports'])}\033[0m):")            
            if grouped:
                print (f"{' '.join(self.group_ip_addresses(cmd_info['hosts_with_open_ports']))}")
            else:
                print (f"{' '.join([str(x) for x in cmd_info['hosts_with_open_ports']])}")
            
            print (f"\n[+] Ports open TCP (\033[92m{len(cmd_info['ports_tcp_open'])}\033[0m):") 
            if grouped:                                       
                print (f"{','.join(self.group_ports(cmd_info['ports_tcp_open']))}")
            else:         
                print (f"{','.join([str(x) for x in cmd_info['ports_tcp_open']])}")        

            print (f"\n[+] Ports open UDP (\033[92m{len(cmd_info['ports_udp_open'])}\033[0m):") 
            if grouped:                                       
                print (f"{','.join(self.group_ports(cmd_info['ports_udp_open']))}")
            else:         
                print (f"{','.join([str(x) for x in cmd_info['ports_udp_open']])}")             
           
    
    # lock object to work thread-safe with print function
    s_print_lock = Lock()

    def s_print(self, value):
        """Thread safe print function"""
        self.s_print_lock.acquire()
        print(value)
        self.s_print_lock.release()    

    def execute(self, nmap_cmd, threads, cache_dir, scans, scans_open):
        "execute nmap_cmd in parallel"
        scans_finished = len(scans) - len(scans_open)
        print (f"[+] Progress: \033[1m\033[92m{scans_finished}/{len(scans)} hosts\033[0m")

        # TODO work in memory instead of tmp

        with tempfile.TemporaryDirectory() as tmp_dir :
            with concurrent.futures.ThreadPoolExecutor(max_workers=threads) as executor:
                try:
                    # init nmap scan
                    for scan in executor.map(self.run_nmap, scans_open, repeat(tmp_dir)):
                        try:
                            if (scan.result.returncode == 0):
                                # update progress
                                scans_finished += 1
                                self.s_print (f'\033[A                             \033[A')
                                self.s_print (f'[+] Progress: \033[1m\033[92m{scans_finished}/{len(scans)} hosts\033[0m')  

                                scan_id = scan.get_id()
                                in_base_path = os.path.join(tmp_dir, scan_id)
                                out_base_path = os.path.join(cache_dir, scan_id)

                                # copy files
                                for file_ext in [".xml", ".txt", ".grep", ".log"]:
                                    shutil.copyfile(in_base_path + file_ext, out_base_path + file_ext)

                            else:                            
                                self.s_print(f"[!] Scan of {scan.ip} failed: Error (Result code: {scan.result.returncode})")                            
                        except Exception as exc:
                            self.s_print(f'[!] Catch inside: {exc}')

                except Exception as exc:
                    self.s_print(f'[!] Catch outside: {exc}')
                    # TODO finish processes when catched, because otherwise there will be zombie processes
            
        # TODO + #FIXME After successful scan the terminal often doesn't show a cursor anymore



def main(cli_args=None):
    nparallel = Nparallel()    

    # Parse args
    args = nparallel.args  
    
    # "nmap" command
    if args.command == "nmap":
        # parse nmap command
        nmap_cmd = NmapCommand(args.nmap_command)

        # resolve IPs
        resolved_ips = nparallel.resolve_to_ip_addresses(args.input_list)  
        cache_dir = nparallel.init_scan_cache()  
        
        # check for (un)finished scans        
        scans, scans_open = nparallel.get_scans (nmap_cmd.get_id(), resolved_ips, args.force_scan)

        # run scans
        start_time = datetime.now()  
        print (f"[*] Cmd id: \033[1m{nmap_cmd.get_id()}\033[0m | Nmap base cmd: \033[1m{nmap_cmd.cmd_str}\033[0m | Threads: \033[1m{args.threads}\033[0m")
        print (f"[*] Start: {start_time.strftime(f'%H:%M:%S')}")      
        nparallel.execute(nmap_cmd, args.threads, cache_dir, scans, scans_open) 
        end_time = datetime.now()
        print (f"[*] End: {end_time.strftime(f'%H:%M:%S')} (\033[1m\033[92mFinished\033[0m in {'{:.2f}'.format((end_time - start_time).total_seconds())} seconds)")
        
        # generate out_files
        export_builder = NmapExportBuilder()

        # set out_all
        if args.out_all:
            args.out_xml = args.out_all + ".xml"
            args.out_normal = args.out_all + ".txt"
            args.out_grepable = args.out_all + ".grep"
        
        # print message if no export file was provided or export files
        if args.out_all == False and \
           args.out_xml == False and \
           args.out_normal == False and \
           args.out_grepable == False and \
           args.out_log == False and \
           args.out_excel == False and \
           args.args == False:
            s_print (f"\nNo output file location provided - run same command again with -oX/-oN/-oG/-oL/-oC/-oE/-oW/-oA option")
        else:
            # leave a little space
            print ()
            if args.out_xml:
                export_builder.export(".xml", cache_dir, nmap_cmd, resolved_ips, start_time, end_time, args.out_xml)
                print (f"[+] XML file saved at '\033[1m{os.path.join(os.getcwd(), args.out_xml)}\033[0m'")
            if args.out_normal:
                export_builder.export(".txt", cache_dir, nmap_cmd, resolved_ips, start_time, end_time, args.out_normal)
                print (f"[+] Normal file saved at '\033[1m{os.path.join(os.getcwd(), args.out_normal)}\033[0m'")
            if args.out_grepable:
                export_builder.export(".grep", cache_dir, nmap_cmd, resolved_ips, start_time, end_time, args.out_grepable)
                print (f"[+] Grepable file saved at '\033[1m{os.path.join(os.getcwd(), args.out_grepable)}\033[0m'")   
            if args.out_log:
                export_builder.export(".log", cache_dir, nmap_cmd, resolved_ips, start_time, end_time, args.out_log)
                print (f"[+] Log file saved at '\033[1m{os.path.join(os.getcwd(), args.out_log)}\033[0m'")    
            if args.out_csv:
                export_builder.export(".csv", cache_dir, nmap_cmd, resolved_ips, start_time, end_time, args.out_csv)
                print (f"[+] CSV file saved at '\033[1m{os.path.join(os.getcwd(), args.out_csv)}\033[0m'")  
            if args.out_excel:
                export_builder.export(".xlsx", cache_dir, nmap_cmd, resolved_ips, start_time, end_time, args.out_excel)
                print (f"[+] XLSX file saved at '\033[1m{os.path.join(os.getcwd(), args.out_excel)}\033[0m'")
            if args.out_word:
                export_builder.export(".docx", cache_dir, nmap_cmd, resolved_ips, start_time, end_time, args.out_word)
                print (f"[+] DOCX file saved at '\033[1m{os.path.join(os.getcwd(), args.out_word)}\033[0m'")
            
    
    # "ls" command
    elif args.command == "ls":
        if args.cmd_id:
            cmd_info = nparallel.get_cmd_info(args.cmd_id)
            nparallel.print_cmd_info(cmd_info, args.grouped)
        else:
            scans = nparallel.get_cache_info()
            nparallel.print_cache_info(scans)

    # "rm" command
    elif args.command == "rm":
        if args.cmd_id:
            nparallel.remove_cache_entries(args.cmd_id)
        else:
            nparallel.remove_cache_entries()

    # leave some space at the end
    print ()

    
if __name__ == '__main__':
    main()

